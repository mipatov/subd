import docx
from fielddict import *
from datetime import datetime

def report(title,filter,tabsrc,path = ''):
    doc = docx.Document()

    doc.add_heading(title, 0)
    if ";" in filter:
        filters = filter.split(";")
        print(filters)
        for fltr in filters:
            if fltr == " ": break
            doc.add_paragraph(fltr)
    else:
        doc.add_paragraph(filter)
    doc.add_paragraph("")

    n, m = len(tabsrc[0]), len(tabsrc)
    table = doc.add_table(rows=m+1, cols=n)
    table.style = 'Table Grid'

    headers = GetTupleOfFullName(tabsrc[0].keys())
    print(headers)
    for col in range(n):
        cell = table.cell(0, col)
        # записываем в ячейку данные
        cell.text = headers[col]

    keys = list(tabsrc[0].keys())
    # заполняем таблицу данными
    for row in range(m):
        for col in range(n):
            # получаем ячейку таблицы
            cell = table.cell(row+1, col)
            # записываем в ячейку данные

            cell.text = str(tabsrc[row][keys[col]])

    now = datetime.now()
    date_time = now.strftime("%d-%m-%Y %H.%M.%S")

    doc.add_paragraph(f'Создан: {date_time}')
    filename = f'{path}/{title} {date_time}.docx'
    doc.save(filename)

    return filename
